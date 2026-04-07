#!/usr/bin/env python3
"""
One-time generator for 5 master DOCX contract templates.
The resulting .docx files are the source of truth (not this script).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'public', 'masters')
os.makedirs(OUTPUT_DIR, exist_ok=True)

DISCLAIMER = (
    "UWAGA: Niniejszy wzór ma charakter wyłącznie informacyjny i edukacyjny. "
    "Przed użyciem należy dostosować go do konkretnej sytuacji oraz "
    "zweryfikować pod kątem prawnym z wykwalifikowanym prawnikiem. "
    "Majster.AI nie ponosi odpowiedzialności za skutki użycia tego wzoru."
)

def add_header(doc, title, subtitle=None):
    """Add document header with title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
    doc.add_paragraph()

def add_section_title(doc, number, title):
    """Add numbered section title."""
    p = doc.add_paragraph()
    run = p.add_run(f"§ {number}. {title}")
    run.bold = True
    run.font.size = Pt(12)

def add_para(doc, text, bold=False):
    """Add a paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_placeholder(doc, label):
    """Add a placeholder field."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}]")
    run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
    run.bold = True
    return p

def add_party_block(doc, party_label):
    """Add a standard party identification block."""
    add_para(doc, f"{party_label}:")
    add_placeholder(doc, "Imię i nazwisko / Nazwa firmy")
    add_placeholder(doc, "Adres siedziby / zamieszkania")
    add_placeholder(doc, "NIP / PESEL")
    add_placeholder(doc, "REGON (jeśli dotyczy)")
    add_placeholder(doc, "Numer telefonu")
    add_placeholder(doc, "Adres e-mail")
    add_placeholder(doc, "Reprezentowany przez (jeśli dotyczy)")
    doc.add_paragraph()

def add_signature_block(doc):
    """Add signature blocks for both parties."""
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "________________________"
    table.cell(0, 1).text = "________________________"
    table.cell(1, 0).text = "Wykonawca"
    table.cell(1, 1).text = "Zamawiający"
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = ""
    table.cell(3, 0).text = "[Data i miejsce]"
    table.cell(3, 1).text = "[Data i miejsce]"

def add_disclaimer(doc):
    """Add legal disclaimer at the end."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(DISCLAIMER)
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.italic = True

print("Part 1 (helpers) loaded OK")
